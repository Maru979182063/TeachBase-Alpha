from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "项目立项设计_v0.1.md"
OUT_DIR = ROOT / "outputs" / "brd_exports"
DOCX_PATH = OUT_DIR / "讲义加工与题目治理项目说明_v0.4.docx"


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(11)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)


def add_image(doc: Document, rel_path: str, alt_text: str) -> None:
    image_path = (ROOT / rel_path).resolve()
    if not image_path.exists():
        return
    image_width = Inches(9.6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=image_width)
    cap = doc.add_paragraph(alt_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1: \2", text)
    return text.strip()


def export_docx() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    first_title_used = False
    for raw in lines:
        line = raw.rstrip()
        if not line:
            continue

        if line.startswith("# "):
            if not first_title_used:
                add_title(doc, clean_inline(line[2:]))
                first_title_used = True
            else:
                doc.add_heading(clean_inline(line[2:]), level=1)
            continue

        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
            continue

        if line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
            continue

        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if img_match:
            add_image(doc, img_match.group(2), img_match.group(1) or "图片")
            continue

        if line.startswith("- "):
            doc.add_paragraph(clean_inline(line[2:]), style="List Bullet")
            continue

        doc.add_paragraph(clean_inline(line))

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = export_docx()
    print(path)
